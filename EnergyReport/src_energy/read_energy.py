import os
import io
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


# =========================
# ✅ PATH SETUP
# =========================
BASE_DIR = os.path.dirname(os.path.abspath(__file__))

def get_image_path(filename):
    path = os.path.join(BASE_DIR, "..", "..", "images", filename)
    if not os.path.exists(path):
        raise FileNotFoundError(f"Image NOT found: {path}")
    return path


# =========================
# ✅ TOC FUNCTION (SAFE)
# =========================
def add_table_of_contents(doc):

    doc.add_page_break()

    # Title
    p = doc.add_paragraph()
    run = p.add_run("Contents")
    run.bold = True
    run.font.size = Pt(14)

    # TOC field
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()

    fldChar_begin = OxmlElement('w:fldChar')
    fldChar_begin.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'

    fldChar_separate = OxmlElement('w:fldChar')
    fldChar_separate.set(qn('w:fldCharType'), 'separate')

    fldChar_end = OxmlElement('w:fldChar')
    fldChar_end.set(qn('w:fldCharType'), 'end')

    run._r.append(fldChar_begin)
    run._r.append(instrText)
    run._r.append(fldChar_separate)
    run._r.append(fldChar_end)


# =========================
# ✅ MAIN FUNCTION
# =========================
def extractReport(deg_0, deg_90, file_180, file_270, proposed_file):

    doc = Document()

    # Margins
    section = doc.sections[0]
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    # =========================
    # ✅ PAGE 1: COVER
    # =========================

    table = doc.add_table(rows=1, cols=2)
    table.autofit = False

    left_cell = table.rows[0].cells[0]
    right_cell = table.rows[0].cells[1]

    # LEFT IMAGE
    run = left_cell.paragraphs[0].add_run()
    run.add_picture(get_image_path("energy_tree.png"), width=Inches(3))

    # RIGHT CONTENT FUNCTION
    def add_right(text, size, bold=False):
        p = right_cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = p.add_run(text)
        r.font.size = Pt(size)
        r.bold = bold
        r.font.color.rgb = RGBColor(138, 139, 58)
        return p

    add_right("Energy", 36, True)
    add_right("Analysis", 36, True)
    add_right("Report", 36, True)

    # Green line
    add_right("________________________", 12)

    right_cell.add_paragraph("")

    # Project name (bigger)
    p = right_cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("MFAR Red Sanders,\nBangalore")
    r.font.size = Pt(16)
    r.bold = True

    right_cell.add_paragraph("")

    # Submitted text (smaller)
    p = right_cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(
        "Submitted by:\n"
        "Environmental Design Solutions Pvt Ltd.,\n"
        "D1/25, Basement, Vasant Vihar,\n"
        "New Delhi, 110057, India."
    )
    r.font.size = Pt(9)

    right_cell.add_paragraph("")

    # Logo
    p = right_cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run()
    run.add_picture(get_image_path("logo.png"), width=Inches(1.5))

    # =========================
    # ✅ PAGE 2: TOC
    # =========================
    add_table_of_contents(doc)

    # =========================
    # ✅ PAGE 3+: CONTENT (for TOC)
    # =========================
    doc.add_page_break()

    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph("Sample content...")

    doc.add_heading("Analysis Methodology", level=1)
    doc.add_paragraph("Sample content...")

    doc.add_heading("Area Statement", level=2)
    doc.add_paragraph("Sample content...")

    # =========================
    # ✅ SAVE (ONLY ONCE)
    # =========================
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)

    return file_stream